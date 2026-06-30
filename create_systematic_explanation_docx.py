#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Create a Word document explaining how to open the web calculator and the statistical methods used."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "validation_outputs_10year"
DOCX_PATH = BASE_DIR / "10年风险网页计算器_打开方法与统计学说明.docx"
THRESHOLD_PATH = BASE_DIR / "risk_thresholds_10_year_Youden_IPCW.csv"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def set_doc_style(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_paragraph(doc, text=""):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return p


def add_table_from_dataframe(doc, df, title=None):
    if title:
        add_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(hdr[i], col, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.4f}"
            else:
                text = val
            set_cell_text(cells[i], text)
    doc.add_paragraph()


def add_picture_if_exists(doc, image_path, caption):
    image_path = Path(image_path)
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.8))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.name = "Times New Roman"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(9)


def main():
    doc = Document()
    set_doc_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("10年风险网页计算器：打开方法与统计学说明")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_paragraph(doc, "本文档用于说明10年死亡风险网页计算器如何打开、页面如何显示、底层风险如何计算、使用了哪些统计方法，以及DCA、IPCW time-dependent ROC Youden、校准曲线和高/低风险分层如何解释。")

    add_heading(doc, "1. 文件位置与网页启动方法", 1)
    add_paragraph(doc, "网页计算器所在文件夹：")
    add_code_block(doc, str(BASE_DIR))
    add_paragraph(doc, "推荐用终端运行以下命令启动网页：")
    add_code_block(doc, f"cd \"{BASE_DIR}\"\nPORT=5001 FLASK_DEBUG=false /Users/gubaichuan/miniconda3/bin/python app.py")
    add_paragraph(doc, "看到终端输出 Running on http://127.0.0.1:5001 后，在浏览器中打开：")
    add_code_block(doc, "http://127.0.0.1:5001")
    add_paragraph(doc, "也可以在另一个终端窗口运行：")
    add_code_block(doc, "open http://127.0.0.1:5001")
    add_paragraph(doc, "如果5001端口被占用，可换成其他端口，例如：")
    add_code_block(doc, f"cd \"{BASE_DIR}\"\nPORT=5022 FLASK_DEBUG=false /Users/gubaichuan/miniconda3/bin/python app.py\nopen http://127.0.0.1:5022")
    add_paragraph(doc, "关闭服务时，在运行网页服务的终端窗口按 Ctrl + C。")

    add_heading(doc, "2. 网页页面显示内容", 1)
    add_paragraph(doc, "网页首页显示两个预测目标：")
    add_bullet(doc, "10-Year All-cause Mortality Risk：10年全因死亡风险。")
    add_bullet(doc, "10-Year Cardiovascular Mortality Risk：10年心血管死亡风险。")
    add_paragraph(doc, "页面输入患者临床指标后，点击 Predict，网页返回：")
    add_bullet(doc, "10年预测风险百分比；")
    add_bullet(doc, "高风险或低/非高风险分层；")
    add_bullet(doc, "使用的高风险cutoff；")
    add_bullet(doc, "SHAP瀑布图，用于解释各变量对预测风险的贡献。")
    add_paragraph(doc, "当前页面已删除Medium Risk，只保留 High Risk 与 Low / Non-high Risk 两层。")

    add_heading(doc, "3. 10年风险的底层计算方法", 1)
    add_paragraph(doc, "本网页计算器使用生存模型，而不是普通二分类模型。输出风险定义为固定时间点累计事件概率：")
    add_code_block(doc, "10年死亡风险 = 1 - S(120 months)")
    add_paragraph(doc, "其中 S(120 months) 为模型估计的120个月生存概率。全因死亡模型和心血管死亡模型均通过 predict_survival_function 计算生存函数，再取 1 - S(120 months)。")
    add_bullet(doc, "全因死亡模型：GradientBoostingSurvivalAnalysis。")
    add_bullet(doc, "心血管死亡模型：RandomSurvivalForest。")
    add_paragraph(doc, "因此，网页显示的风险可以称为10-year mortality probability，而不是risk score或普通分类概率。")

    add_heading(doc, "4. 结局定义与数据来源", 1)
    add_paragraph(doc, "用于生成统计图和阈值的脚本为：")
    add_code_block(doc, str(BASE_DIR / "generate_10year_validation_figures.py"))
    add_paragraph(doc, "脚本读取训练集和测试集标准化数据，仅用于分析；所有输出文件均保存在网页计算器文件夹内。")
    add_bullet(doc, "训练集用于推导高风险阈值。")
    add_bullet(doc, "测试集用于验证ROC、DCA和校准曲线。")
    add_bullet(doc, "全因死亡事件定义：mortstat == 1。")
    add_bullet(doc, "心血管死亡事件定义：mortstat == 1 且 ucod_leading in {1, 5}。")
    add_bullet(doc, "随访时间变量：permth_exm，单位为月。")

    add_heading(doc, "5. 高风险阈值：IPCW time-dependent ROC Youden", 1)
    add_paragraph(doc, "高风险阈值通过训练集中的10年IPCW time-dependent ROC曲线确定。Youden指数定义为：")
    add_code_block(doc, "Youden index = sensitivity + specificity - 1")
    add_paragraph(doc, "取Youden指数最大的预测风险值作为高风险cutoff。IPCW（inverse probability of censoring weighting）用于处理10年前删失造成的偏倚。")
    add_paragraph(doc, "重要原则：阈值应在训练集/derivation cohort 中推导，测试集/validation cohort 用于验证，不应在测试集中重新选择cutoff。")

    thresholds = pd.read_csv(THRESHOLD_PATH)
    table_df = pd.DataFrame({
        "结局": thresholds["outcome_label"],
        "模型": thresholds["model"],
        "训练集IPCW AUC": thresholds["auc_ipcw_derivation"],
        "高风险阈值": thresholds["threshold_ipcw"].map(lambda x: f"{x:.2%}"),
        "敏感度": thresholds["sensitivity_ipcw_derivation"],
        "特异度": thresholds["specificity_ipcw_derivation"],
        "测试集IPCW AUC": thresholds["auc_ipcw_validation"],
    })
    add_table_from_dataframe(doc, table_df, "当前训练集推导的10年IPCW Youden阈值如下：")

    add_heading(doc, "6. 当前网页风险分层规则", 1)
    add_paragraph(doc, "网页计算器采用二分类风险分层：")
    add_bullet(doc, "High Risk：预测10年风险 ≥ 训练集推导的IPCW Youden cutoff。")
    add_bullet(doc, "Low / Non-high Risk：预测10年风险 < 训练集推导的IPCW Youden cutoff。")
    add_paragraph(doc, "具体阈值为：")
    add_bullet(doc, "全因死亡：High Risk ≥ 7.75%；Low / Non-high Risk < 7.75%。")
    add_bullet(doc, "心血管死亡：High Risk ≥ 5.47%；Low / Non-high Risk < 5.47%。")
    add_paragraph(doc, "不再使用低/中/高三分类，原因是中风险和低风险阈值若采用“高风险阈值的一半”，缺乏独立统计学依据；二分类更容易解释，也更适合论文和审稿回复。")

    add_heading(doc, "7. DCA曲线说明", 1)
    add_paragraph(doc, "DCA（decision curve analysis）用于评估模型在不同阈值概率下是否具有临床净获益。DCA图包含三条策略线：")
    add_bullet(doc, "Model：根据模型预测风险决定是否判为高风险。")
    add_bullet(doc, "Treat all：所有人均视为高风险。")
    add_bullet(doc, "Treat none：所有人均不干预。")
    add_paragraph(doc, "如果模型曲线在某个阈值范围内高于Treat all和Treat none，则说明模型在该阈值范围内有更高净获益。")
    add_paragraph(doc, "本项目DCA图横轴展示0%–30%的阈值概率。心血管死亡事件率较低，net benefit本身较小，因此作图时裁剪了较大的负向Treat-all区域，以便显示0以上的临床相关净获益区间。")
    add_picture_if_exists(doc, OUTPUT_DIR / "DCA_10year_all_cause_IPCW.png", "图1. 10年全因死亡DCA曲线")
    add_picture_if_exists(doc, OUTPUT_DIR / "DCA_10year_cardiovascular_IPCW.png", "图2. 10年心血管死亡DCA曲线")

    add_heading(doc, "8. ROC曲线说明", 1)
    add_paragraph(doc, "10年IPCW time-dependent ROC曲线用于评价模型在120个月固定时间点区分事件与非事件的能力，并用于展示AUC。为避免误解，测试集和外部验证集ROC图不标注Youden红点，因为cutoff是在训练集中推导的，不应在测试集或外部验证集中重新选择。训练集cutoff应用到测试集/外部验证集后的敏感度和特异度应放在表格或文字中报告。")
    add_picture_if_exists(doc, OUTPUT_DIR / "ROC_10year_all_cause_IPCW.png", "图3. 内部测试集10年全因死亡IPCW time-dependent ROC曲线（不标注Youden红点）")
    add_picture_if_exists(doc, OUTPUT_DIR / "ROC_10year_cardiovascular_IPCW.png", "图4. 内部测试集10年心血管死亡IPCW time-dependent ROC曲线（不标注Youden红点）")

    add_heading(doc, "9. 校准曲线说明", 1)
    add_paragraph(doc, "校准曲线用于比较模型预测的10年风险与观察到的10年事件率是否一致。理想情况下，点应接近45度参考线。Brier score用于量化概率预测误差，数值越低通常表示整体预测误差越小。")
    add_picture_if_exists(doc, OUTPUT_DIR / "Calibration_10year_all_cause.png", "图5. 10年全因死亡校准曲线")
    add_picture_if_exists(doc, OUTPUT_DIR / "Calibration_10year_cardiovascular.png", "图6. 10年心血管死亡校准曲线")

    add_heading(doc, "10. CHARLS外部验证补充分析", 1)
    external_dir = BASE_DIR / "external_charls_outputs_10year"
    external_summary_path = external_dir / "external_charls_10year_ipcw_validation_summary.csv"
    if external_summary_path.exists():
        ext = pd.read_csv(external_summary_path)
        ext_table = pd.DataFrame({
            "结局": ext["outcome_label"],
            "模型": ext["model"],
            "外部IPCW AUC": ext["auc_ipcw_external"],
            "训练集cutoff": ext["training_derived_cutoff"].map(lambda x: f"{x:.2%}"),
            "外部Sensitivity": ext["sensitivity_at_training_cutoff"],
            "外部Specificity": ext["specificity_at_training_cutoff"],
            "外部Brier score": ext["brier_score_ipcw_external"],
        })
        add_paragraph(doc, "CHARLS外部验证集仅用于验证，不重新训练模型，不重新推导cutoff。外部验证预测值同样采用网页一致的10年风险定义：1 - S(120 months)。外部ROC图不标注Youden红点；DCA竖线表示训练集推导cutoff。")
        add_table_from_dataframe(doc, ext_table, "CHARLS外部10年IPCW验证结果：")
        add_picture_if_exists(doc, external_dir / "External_CHARLS_ROC_10year_all_cause_IPCW.png", "图7. CHARLS外部10年全因死亡IPCW ROC曲线（不标注Youden红点）")
        add_picture_if_exists(doc, external_dir / "External_CHARLS_ROC_10year_cardiovascular_IPCW.png", "图8. CHARLS外部10年心血管死亡IPCW ROC曲线（不标注Youden红点）")
        add_picture_if_exists(doc, external_dir / "External_CHARLS_Calibration_10year_all_cause.png", "图9. CHARLS外部10年全因死亡校准曲线")
        add_picture_if_exists(doc, external_dir / "External_CHARLS_Calibration_10year_cardiovascular.png", "图10. CHARLS外部10年心血管死亡校准曲线")
        add_picture_if_exists(doc, external_dir / "External_CHARLS_DCA_10year_all_cause_IPCW.png", "图11. CHARLS外部10年全因死亡DCA曲线")
        add_picture_if_exists(doc, external_dir / "External_CHARLS_DCA_10year_cardiovascular_IPCW.png", "图12. CHARLS外部10年心血管死亡DCA曲线")
        add_paragraph(doc, "需要特别注意：CHARLS外部AUC表现较好，但外部校准曲线提示绝对风险可能存在明显过度预测，尤其是全因死亡。这说明模型具有一定排序/区分能力，但跨队列绝对风险概率需要谨慎解释，必要时可考虑外部再校准。")

    add_heading(doc, "11. 输出文件清单", 1)
    add_paragraph(doc, "主要内部统计输出文件位于：")
    add_code_block(doc, str(OUTPUT_DIR))
    add_bullet(doc, "risk_thresholds_10_year_Youden_IPCW.csv：训练集推导阈值和测试集验证结果。")
    add_bullet(doc, "ROC_10year_*_IPCW.pdf/png：内部测试集10年IPCW time-dependent ROC图。")
    add_bullet(doc, "DCA_10year_*_IPCW.pdf/png：内部测试集10年DCA曲线。")
    add_bullet(doc, "Calibration_10year_*.pdf/png：内部测试集10年校准曲线。")
    add_bullet(doc, "dca_net_benefit_10year_*.csv：内部测试集DCA净获益表。")
    add_bullet(doc, "validation_summary_10year.txt：内部验证中文汇总。")
    add_paragraph(doc, "CHARLS外部验证输出文件位于：")
    add_code_block(doc, str(BASE_DIR / "external_charls_outputs_10year"))
    add_bullet(doc, "External_CHARLS_ROC_10year_*_IPCW.pdf/png：CHARLS外部10年IPCW ROC图。")
    add_bullet(doc, "External_CHARLS_Calibration_10year_*.pdf/png：CHARLS外部10年校准曲线。")
    add_bullet(doc, "External_CHARLS_DCA_10year_*_IPCW.pdf/png：CHARLS外部10年DCA曲线。")
    add_bullet(doc, "external_charls_10year_ipcw_validation_summary.csv/txt：CHARLS外部验证汇总。")

    add_heading(doc, "12. 推荐论文或回复审稿人表述", 1)
    add_paragraph(doc, "英文表述：")
    add_paragraph(doc, "The predicted 10-year mortality risk was calculated as one minus the estimated survival probability at 120 months. The high-risk threshold was derived in the training cohort using the Youden index from an inverse probability of censoring weighted time-dependent ROC analysis at 10 years. Model performance and clinical utility were then evaluated in the test cohort using time-dependent ROC analysis, calibration curves, and decision curve analysis. Participants were classified as high risk if their predicted 10-year risk was greater than or equal to the training-derived Youden threshold, and as low/non-high risk otherwise.")
    add_paragraph(doc, "中文解释：")
    add_paragraph(doc, "10年预测风险定义为120个月时的累计事件概率，即1减去120个月生存概率。高风险阈值在训练集中通过10年IPCW时间依赖ROC曲线的Youden指数确定。随后在测试集中通过时间依赖ROC、校准曲线和DCA评估模型性能及临床净获益。网页计算器中，预测10年风险大于或等于训练集推导Youden阈值者定义为高风险，其余定义为低风险或非高风险。")
    add_paragraph(doc, "保守补充：风险分层主要用于风险沟通和研究展示，不应被解释为直接治疗决策阈值。")

    add_heading(doc, "13. 注意事项", 1)
    add_bullet(doc, "网页计算器为研究工具，输出结果不应直接替代临床判断。")
    add_bullet(doc, "高风险阈值是统计学阈值，不等同于治疗启动阈值。")
    add_bullet(doc, "若未来更换模型或数据，应重新生成IPCW Youden阈值、DCA和校准曲线。")
    add_bullet(doc, "若部署到服务器，需确认部署包中包含 risk_thresholds_10_year_Youden_IPCW.csv。")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
